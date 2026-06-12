"""
Builds a .docx resume in the JobStreet PH "professional template" layout
from the plain text output produced by the rewrite prompt.

Expected input format (produced by SYSTEM_PROMPT):

    Full Name | Target Title | email@example.com | +63 900 000 0000 | City, PH

    SUMMARY
    ...paragraph...

    SKILLS
    • Skill one
    • Skill two

    EXPERIENCE
    Job Title
    Company Name
    Month YYYY to Month YYYY
    • bullet
    • bullet

    Job Title 2
    Company Name 2
    Month YYYY to Month YYYY
    • bullet

    EDUCATION
    Degree or Certificate
    Institution Name
    Year

    PROJECTS
    Project Name, Role or Tech Stack
    • bullet
    • bullet

    CERTIFICATIONS
    • Certification one

Layout produced (mirrors the JobStreet two column template):

    [optional photo]  Full Name
                       Target Title
                       email | phone | location
    --------------------------------------------------
    | SKILLS          | SUMMARY                       |
    | • ...           | ...                           |
    |                  | EXPERIENCE                    |
    | EDUCATION        | Job Title                     |
    | Degree           | Company, Dates                |
    | Institution      | • bullet                      |
    | Year             |                                |
    |                  | PROJECTS                       |
    | CERTIFICATIONS   | Project Name                  |
    | • ...            | • bullet                       |
    --------------------------------------------------

Only sections present in the input text are rendered.
"""

import io
import re

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Theme (kept in sync with the frontend's eye comfortable palette)
# ---------------------------------------------------------------------------
COLOR_PRIMARY = RGBColor(0x3F, 0x5E, 0x4E)   # muted forest green - headings
COLOR_TEXT = RGBColor(0x2B, 0x2B, 0x28)      # near black, warm
COLOR_MUTED = RGBColor(0x70, 0x75, 0x6B)     # muted gray green - dates/meta
COLOR_RULE = RGBColor(0xB8, 0xC4, 0xB6)      # light sage - section underline

KNOWN_SECTIONS = [
    "SUMMARY",
    "SKILLS",
    "EXPERIENCE",
    "EDUCATION",
    "PROJECTS",
    "CERTIFICATIONS",
    "INTERESTS",
]

# Sections rendered in the narrow left sidebar; everything else goes right.
SIDEBAR_SECTIONS = {"SKILLS", "EDUCATION", "CERTIFICATIONS", "INTERESTS"}


# ---------------------------------------------------------------------------
# Parsing
# ---------------------------------------------------------------------------
def parse_resume_text(text: str) -> dict:
    """
    Splits the rewritten resume text into a structured dict:
        {
          "name": str,
          "title": str | None,
          "contact": [str, ...],
          "sections": { "SUMMARY": "...", "EXPERIENCE": [ {...}, ... ], ... }
        }
    """
    lines = [l.rstrip() for l in text.strip().splitlines()]

    # --- Header line: "Name | Title | email | phone | location"
    header_line = ""
    body_start = 0
    for i, line in enumerate(lines):
        if line.strip():
            header_line = line.strip()
            body_start = i + 1
            break

    parts = [p.strip() for p in header_line.split("|") if p.strip()]
    name = parts[0] if parts else "Your Name"
    title = parts[1] if len(parts) > 1 else None
    contact = parts[2:] if len(parts) > 2 else parts[1:] if len(parts) > 1 and title is None else parts[2:]
    if len(parts) == 2:
        # Ambiguous: could be "Name | Title" or "Name | contact". Heuristic:
        # treat as title only (more common for resumes); contact stays empty.
        contact = []
    elif len(parts) > 2:
        contact = parts[2:]

    # --- Walk remaining lines, grouping into sections
    sections: dict = {}
    current_section = None
    buffer: list[str] = []

    def flush():
        nonlocal buffer, current_section
        if current_section is None:
            buffer = []
            return
        raw = "\n".join(buffer).strip("\n")
        if current_section == "SUMMARY":
            sections[current_section] = raw.strip()
        elif current_section in ("SKILLS", "CERTIFICATIONS", "INTERESTS"):
            items = []
            for l in buffer:
                l = l.strip()
                if not l:
                    continue
                l = re.sub(r"^[•\-*]\s*", "", l)
                items.append(l)
            sections[current_section] = items
        elif current_section in ("EXPERIENCE", "PROJECTS"):
            sections[current_section] = _parse_entries(buffer)
        elif current_section == "EDUCATION":
            sections[current_section] = _parse_education(buffer)
        buffer = []

    for line in lines[body_start:]:
        stripped = line.strip()

        # Clean up markdown formatting (bolding, hashes) and colons
        cleaned = re.sub(r'^#+\s*|\*+|_|:$', '', stripped).strip()
        upper = cleaned.upper()

        # Removed the strict `(stripped == upper)` condition 
        # to make the parser resilient against LLM formatting quirks
        if upper in KNOWN_SECTIONS:
            flush()
            current_section = upper
            continue
        buffer.append(line)

    flush()

    return {
        "name": name,
        "title": title,
        "contact": contact,
        "sections": sections,
    }


def _parse_entries(buffer: list[str]) -> list[dict]:
    """
    Parses EXPERIENCE / PROJECTS blocks. Each entry is separated by a blank
    line. Within an entry, leading non-bullet lines are "meta" (title,
    company, dates) and bullet lines (•, -, *) are details.
    """
    entries = []
    current: dict | None = None

    for raw in buffer:
        line = raw.strip()
        if not line:
            if current and (current["meta"] or current["bullets"]):
                entries.append(current)
                current = None
            continue

        if current is None:
            current = {"meta": [], "bullets": []}

        if re.match(r"^[•\-*]\s*", line):
            current["bullets"].append(re.sub(r"^[•\-*]\s*", "", line))
        else:
            current["meta"].append(line)

    if current and (current["meta"] or current["bullets"]):
        entries.append(current)

    return entries


def _parse_education(buffer: list[str]) -> list[dict]:
    """
    Parses EDUCATION block into entries. Groups are separated by blank
    lines; each group's lines become "meta" lines (degree, institution,
    year). Bullets (rare for education) are kept as "bullets".
    """
    entries = []
    current: dict | None = None

    for raw in buffer:
        line = raw.strip()
        if not line:
            if current and current["meta"]:
                entries.append(current)
                current = None
            continue

        if current is None:
            current = {"meta": [], "bullets": []}

        if re.match(r"^[•\-*]\s*", line):
            current["bullets"].append(re.sub(r"^[•\-*]\s*", "", line))
        else:
            current["meta"].append(line)

    if current and current["meta"]:
        entries.append(current)

    return entries


# ---------------------------------------------------------------------------
# Docx helpers
# ---------------------------------------------------------------------------
def _set_cell_background(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tblPr.append(borders)


def _set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tcPr.append(mar)


def _add_bottom_border(paragraph, color="B8C4B6", size=6):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _heading(container, text: str, size=12):
    p = container.add_paragraph()
    run = p.add_run(text.title())
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = COLOR_PRIMARY
    p.space_before = Pt(10)
    p.space_after = Pt(2)
    _add_bottom_border(p)
    return p


def _bullet(container, text: str, size=10):
    p = container.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = COLOR_TEXT
    p.paragraph_format.space_after = Pt(2)
    return p


def _body(container, text: str, size=10, italic=False, color=None, bold=False, space_after=4):
    p = container.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    run.bold = bold
    run.font.color.rgb = color or COLOR_TEXT
    p.paragraph_format.space_after = Pt(space_after)
    return p


# ---------------------------------------------------------------------------
# Section renderers
# ---------------------------------------------------------------------------
def _render_skills(container, items: list[str]):
    _heading(container, "Skills")
    for item in items:
        _bullet(container, item)


def _render_education(container, entries: list[dict]):
    _heading(container, "Education")
    for entry in entries:
        meta = entry["meta"]
        if meta:
            _body(container, meta[0], bold=True, size=10, space_after=0)
        for line in meta[1:]:
            _body(container, line, size=9, color=COLOR_MUTED, space_after=0)
        for b in entry["bullets"]:
            _bullet(container, b, size=9)
        _body(container, "", size=4)  # spacer


def _render_certifications(container, items: list[str]):
    _heading(container, "Certifications")
    for item in items:
        _bullet(container, item)


def _render_interests(container, items: list[str]):
    _heading(container, "Interests")
    for item in items:
        _bullet(container, item)


def _render_summary(container, text: str):
    _heading(container, "Summary")
    _body(container, text, size=10)


def _render_experience_like(container, heading: str, entries: list[dict]):
    _heading(container, heading)
    for entry in entries:
        meta = entry["meta"]
        if meta:
            _body(container, meta[0], bold=True, size=11, space_after=0)
        if len(meta) > 1:
            _body(container, " | ".join(meta[1:]), size=9, color=COLOR_MUTED, italic=True, space_after=2)
        for b in entry["bullets"]:
            _bullet(container, b)
        _body(container, "", size=4)  # spacer between entries


SECTION_RENDERERS = {
    "SUMMARY": ("right", _render_summary),
    "SKILLS": ("left", _render_skills),
    "EDUCATION": ("left", _render_education),
    "CERTIFICATIONS": ("left", _render_certifications),
    "INTERESTS": ("left", _render_interests),
    "EXPERIENCE": ("right", lambda c, v: _render_experience_like(c, "Experience", v)),
    "PROJECTS": ("right", lambda c, v: _render_experience_like(c, "Projects", v)),
}


# ---------------------------------------------------------------------------
# Main builder
# ---------------------------------------------------------------------------
def build_resume_docx(resume_text: str, photo_bytes: bytes | None = None) -> bytes:
    parsed = parse_resume_text(resume_text)
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.27)   # A4
    section.page_height = Inches(11.69)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.font.color.rgb = COLOR_TEXT

    # --- Header: photo (optional) + name/title/contact ---------------------
    if photo_bytes:
        header_table = doc.add_table(rows=1, cols=2)
        _remove_table_borders(header_table)
        header_table.columns[0].width = Inches(1.3)
        header_table.columns[1].width = Inches(5.77)
        photo_cell, info_cell = header_table.rows[0].cells
        _set_cell_margins(photo_cell, right=120)
        photo_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = photo_cell.paragraphs[0]
        run = p.add_run()
        run.add_picture(io.BytesIO(photo_bytes), width=Inches(1.15))
    else:
        info_cell = doc

    name_p = info_cell.add_paragraph() if photo_bytes else doc.add_paragraph()
    run = name_p.add_run(parsed["name"])
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = COLOR_PRIMARY
    name_p.paragraph_format.space_after = Pt(0)

    if parsed["title"]:
        title_p = info_cell.add_paragraph() if photo_bytes else doc.add_paragraph()
        run = title_p.add_run(parsed["title"])
        run.font.size = Pt(12)
        run.font.color.rgb = COLOR_MUTED
        title_p.paragraph_format.space_after = Pt(2)

    if parsed["contact"]:
        contact_p = info_cell.add_paragraph() if photo_bytes else doc.add_paragraph()
        run = contact_p.add_run("   |   ".join(parsed["contact"]))
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_MUTED
        contact_p.paragraph_format.space_after = Pt(6)

    rule_p = doc.add_paragraph()
    _add_bottom_border(rule_p, color="3F5E4E", size=12)
    rule_p.paragraph_format.space_after = Pt(8)

    # --- Two column body ----------------------------------------------------
    sections = parsed["sections"]
    body_table = doc.add_table(rows=1, cols=2)
    _remove_table_borders(body_table)
    body_table.columns[0].width = Inches(2.3)
    body_table.columns[1].width = Inches(4.77)
    left_cell, right_cell = body_table.rows[0].cells
    _set_cell_margins(left_cell, right=180)
    _set_cell_margins(right_cell, left=180)

    # Render sections in a sensible fixed order, skipping missing ones.
    order = ["SUMMARY", "EXPERIENCE", "PROJECTS", "SKILLS", "EDUCATION", "CERTIFICATIONS", "INTERESTS"]
    for sec_name in order:
        if sec_name not in sections:
            continue
        value = sections[sec_name]
        side, renderer = SECTION_RENDERERS[sec_name]
        container = left_cell if side == "left" else right_cell
        renderer(container, value)

    # Clean up the default empty paragraph python-docx adds to each cell
    for cell in (left_cell, right_cell):
        first = cell.paragraphs[0]
        if not first.text and len(cell.paragraphs) > 1:
            first._p.getparent().remove(first._p)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
